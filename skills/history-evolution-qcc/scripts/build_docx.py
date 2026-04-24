"""
build_docx.py · 7 章 + 附录 DOCX/PDF 生成器（V1.1.3）
=======================================================
V1.1.3 与 V1.0.1 差异：
  - §2 注册资本 / §3 名称 / §6 地址 / §4.3 股东进出 / §5.4 历任董监高 / §附录 大事年表
    全部改为嵌入时间轴 PNG（QCC Advanced Report 风格）
  - §7 总结 改为 Decision Pack 风格摘要卡（借鉴 kyb-opus）
  - 其余章节（§1 概况 / §4.1 当前股东 / §5.2 分支许可 / §5.5 关联方 / §5.6 风险 / §7 财务）保留表格

原则：
  - MCP 返回空 → 整节不输出
  - 每个字段自动附来源标签
"""

from __future__ import annotations

import argparse
import io
import json
import os
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
from mcp_orchestrator import CompanySnapshot, should_render_section  # noqa: E402
from tier_detector import TIER_FINANCING, TIER_LISTED, TIER_MICRO  # noqa: E402
import build_timeline  # noqa: E402


TODAY = date.today().isoformat()

TOP_N_BY_TIER = {
    TIER_MICRO: 5,
    TIER_FINANCING: 10,
    TIER_LISTED: 20,
}

ADVANCED_BLUE = RGBColor(0x1A, 0x73, 0xE8)


def _is_empty_response(data: Any) -> bool:
    if data is None:
        return True
    if isinstance(data, (list, str, tuple)) and not data:
        return True
    if isinstance(data, dict):
        if data.get("_error"):
            return True
        search_result = data.get("搜索结果") or ""
        if any(k in search_result for k in ["未发现", "未找到", "暂无"]):
            return True
        meta_keys = {"企业名称", "摘要", "搜索结果", "_error", "人员名称"}
        content_keys = [k for k in data.keys() if k not in meta_keys]
        if not content_keys:
            return True
        if all(
            not data[k] or (isinstance(data[k], (list, dict)) and not data[k])
            for k in content_keys
        ):
            return True
    return False


def _safe_list(data: Any, key: str) -> list:
    if not isinstance(data, dict):
        return []
    v = data.get(key)
    return v if isinstance(v, list) else []


def _set_default_font(doc: Document, font: str = "Microsoft YaHei") -> None:
    for sty in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        try:
            s = doc.styles[sty]
            s.font.name = font
            s.element.rPr.rFonts.set(qn("w:eastAsia"), font)
        except (KeyError, AttributeError):
            pass


# ---------------------------------------------------------------------------
# V1.1.2：kyb-opus 风格样式工具
# ---------------------------------------------------------------------------

from docx.oxml import OxmlElement

# kyb-opus 配色
KYB_BLUE = "2563EB"           # 主蓝（表头背景、标题文字）
KYB_LIGHT_BG = "EEF4FB"       # 浅蓝（标题/条纹行背景）
KYB_ZEBRA_BG = "F6F8FB"       # 超浅蓝（数据行条纹）


def _set_shading(element_pr, fill_hex: str) -> None:
    """给 paragraph/cell 的 pPr/tcPr 设置 shading。fill_hex='EEF4FB'（无 #）"""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element_pr.append(shd)


def _set_cell_bg(cell, fill_hex: str) -> None:
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    # 去除已有 shd
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    _set_shading(tc_pr, fill_hex)


def _set_para_left_bar(paragraph, color_hex: str = KYB_BLUE, size: int = 32) -> None:
    """给段落左侧加蓝色竖条（kyb-opus 章节标题样式）
    size: 1/8 pt（36 = 4.5pt 粗竖条）
    """
    p_pr = paragraph._p.get_or_add_pPr()
    # 去除已有 border
    for existing in p_pr.findall(qn("w:pBdr")):
        p_pr.remove(existing)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color_hex)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def _set_cell_border(cell, color_hex: str = "D1D5DB") -> None:
    """给单元格加浅灰细边框"""
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """kyb-opus 风格：左侧蓝色竖条 + 蓝色粗字，无背景。
    level=1：大标题（4.5pt 粗竖条 + 16pt bold 蓝字）
    level=2：二级标题（3pt 竖条 + 13pt bold 蓝字）
    level>=3：三级（11pt bold 蓝字，无竖条）
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.bold = True
    run.font.color.rgb = ADVANCED_BLUE

    if level == 1:
        run.font.size = Pt(16)
        _set_para_left_bar(p, KYB_BLUE, size=36)  # 4.5pt 粗竖条
    elif level == 2:
        run.font.size = Pt(13)
        _set_para_left_bar(p, KYB_BLUE, size=24)  # 3pt 竖条
    else:
        run.font.size = Pt(11)


def _add_para(
    doc: Document,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size: int = 10,
    source_tag: str | None = None,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if source_tag:
        tag_run = p.add_run(f"  [来源: {source_tag} · snapshot={TODAY}]")
        tag_run.font.size = Pt(8)
        tag_run.italic = True
        tag_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """kyb-opus 风格表格：实心蓝底白字表头 + 浅蓝/白交替行 + 浅灰细边框"""
    if not rows:
        return
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头行：蓝色背景 + 白色粗体
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, KYB_BLUE)
        _set_cell_border(cell, KYB_BLUE)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.name = "Microsoft YaHei"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 数据行：浅蓝/白交替
    for ridx, row in enumerate(rows, start=1):
        zebra = ridx % 2 == 1  # 奇数行加浅色
        for cidx, val in enumerate(row):
            cell = t.rows[ridx].cells[cidx]
            cell.text = str(val) if val is not None else ""
            if zebra:
                _set_cell_bg(cell, KYB_ZEBRA_BG)
            _set_cell_border(cell, "E5E7EB")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Microsoft YaHei"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)


def _embed_timeline_png(doc: Document, png_bytes: bytes, width_cm: float = 16) -> None:
    """把 PNG bytes 作为图片嵌入文档"""
    if not png_bytes:
        return
    stream = io.BytesIO(png_bytes)
    doc.add_picture(stream, width=Cm(width_cm))


# ---------------------------------------------------------------------------
# 章节渲染
# ---------------------------------------------------------------------------


def render_title(doc: Document, snap: CompanySnapshot) -> None:
    profile = snap.profile
    name = profile.get("企业名称") or snap.keyword
    doc.add_heading(f"{name} · 企业历史沿革与发展历程报告", level=0)
    tier_cn = {"micro": "小微档", "financing": "融资档", "listed": "上市档"}.get(snap.tier, snap.tier)
    _add_para(doc, f"档位：{tier_cn} · 识别依据：{snap.tier_reason}", italic=False, size=9)
    _add_para(
        doc,
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')} · "
        f"调用成本：{snap.cost_summary.get('total_cent', 0)}c "
        f"(¥{snap.cost_summary.get('total_rmb', 0)})",
        italic=False,
        size=9,
    )


def render_s1_overview(doc: Document, snap: CompanySnapshot) -> None:
    profile = snap.profile
    reg = snap.raw.get("mcp__qcc-company__get_company_registration_info") or {}
    tax = snap.raw.get("mcp__qcc-company__get_tax_invoice_info") or {}

    if _is_empty_response(profile) and _is_empty_response(reg):
        return

    _add_heading(doc, "一、公司概况", level=1)

    pairs = [
        ("企业名称", profile.get("企业名称")),
        ("统一社会信用代码", reg.get("统一社会信用代码")),
        ("成立日期", reg.get("成立日期")),
        ("注册资本", reg.get("注册资本")),
        ("实缴资本", reg.get("实缴资本")),
        ("经营期限", reg.get("经营期限") or reg.get("营业期限")),
        ("企业类型", reg.get("企业类型")),
        ("行业", reg.get("国民经济行业") or profile.get("行业")),
        ("经营状态", reg.get("登记状态")),
        ("登记机关", reg.get("登记机关")),
        ("纳税人识别号", tax.get("纳税人识别号")),
    ]
    rows = [[k, str(v)] for k, v in pairs if v]
    if rows:
        _add_table(doc, ["字段", "值"], rows)
        _add_para(doc, "", source_tag="qcc-company.get_company_profile + get_company_registration_info")

    brief = profile.get("简介")
    if brief:
        _add_para(doc, "企业简介：", bold=True)
        _add_para(doc, brief, size=10)


def render_s2_capital(doc: Document, snap: CompanySnapshot) -> None:
    hist_reg = snap.raw.get("mcp__qcc-history__get_historical_registration") or {}
    capi_list = _safe_list(hist_reg, "历史注册资本列表")
    if not capi_list:
        return

    _add_heading(doc, "二、注册资本变更轨迹", level=1)
    png = build_timeline.timeline_capital(capi_list)
    _embed_timeline_png(doc, png, width_cm=15)
    _add_para(
        doc,
        f"共计 {len(capi_list)} 个注册资本节点。",
        size=9,
        italic=False,
        source_tag="qcc-history.get_historical_registration",
    )


def render_s3_name(doc: Document, snap: CompanySnapshot) -> None:
    hist_reg = snap.raw.get("mcp__qcc-history__get_historical_registration") or {}
    name_list = _safe_list(hist_reg, "历史名称列表")
    if not name_list:
        return

    _add_heading(doc, "三、名称变更时间线", level=1)
    png = build_timeline.timeline_name(name_list)
    _embed_timeline_png(doc, png, width_cm=16)
    _add_para(doc, "", size=9, source_tag="qcc-history.get_historical_registration")


def render_s4_equity(doc: Document, snap: CompanySnapshot) -> None:
    current = snap.raw.get("mcp__qcc-company__get_shareholder_info") or {}
    partners = _safe_list(current, "股东信息")
    hist_data = snap.raw.get("mcp__qcc-history__get_historical_shareholders") or {}
    hist_partners = _safe_list(hist_data, "历史股东信息")
    hist_legal = snap.raw.get("mcp__qcc-history__get_historical_legal_rep") or {}
    legal_reps = _safe_list(hist_legal, "历史法定代表人信息")

    if not (partners or hist_partners or legal_reps):
        return

    _add_heading(doc, "四、股权演变", level=1)

    # §4.1 当前股东（表格）
    if partners:
        _add_heading(doc, "4.1 当前股东与出资结构", level=2)
        has_sub_cap = any(p.get("认缴出资额") for p in partners)
        has_shares = any(p.get("持股数") for p in partners)

        if has_shares and not has_sub_cap:
            headers = ["股东名称", "持股比例", "持股数"]
            rows = [[p.get("股东名称") or "", p.get("持股比例") or "", p.get("持股数") or ""] for p in partners]
        elif has_sub_cap:
            headers = ["股东名称", "持股比例", "认缴出资额", "认缴日期", "实缴出资额", "实缴日期"]
            rows = [
                [
                    p.get("股东名称") or "",
                    p.get("持股比例") or "",
                    p.get("认缴出资额") or "",
                    p.get("认缴出资日期") or "",
                    p.get("实缴出资额") or "",
                    p.get("实缴出资日期") or "",
                ]
                for p in partners
            ]
        else:
            headers = ["股东名称", "持股比例"]
            rows = [[p.get("股东名称") or "", p.get("持股比例") or ""] for p in partners]

        _add_table(doc, headers, rows)
        _add_para(doc, "", source_tag="qcc-company.get_shareholder_info")

    # §4.2 初期股东结构变动
    if hist_partners or legal_reps:
        earliest_exits = sorted(
            [p for p in hist_partners if p.get("退出日期")],
            key=lambda p: p.get("退出日期") or "",
        )[:5]
        if earliest_exits or legal_reps:
            _add_heading(doc, "4.2 初期股东结构变动", level=2)
            if earliest_exits:
                _add_para(doc, "早期退出股东（按退出日期升序，最多呈现 5 位）：", bold=True)
                rows = [
                    [
                        p.get("退出日期") or "",
                        p.get("股东名称") or "",
                        p.get("股东类型") or "",
                        p.get("退出时持股比例") or "",
                        p.get("认缴出资额") or "",
                    ]
                    for p in earliest_exits
                ]
                _add_table(doc, ["退出日期", "股东名称", "股东类型", "退出时持股比例", "认缴出资额"], rows)
            if legal_reps:
                _add_para(doc, "历任法定代表人：", bold=True)
                png = build_timeline.timeline_legal_rep(legal_reps)
                _embed_timeline_png(doc, png, width_cm=14)
            _add_para(
                doc,
                "以上记录反映企业设立及早期阶段的股东与法人变动。",
                size=9,
                italic=False,
                source_tag="qcc-history.get_historical_shareholders + get_historical_legal_rep",
            )

    # §4.3 股东进出历史 —— 时间轴
    if hist_partners:
        _add_heading(doc, "4.3 股东进出历史", level=2)
        png = build_timeline.timeline_shareholder_exits(hist_partners)
        _embed_timeline_png(doc, png, width_cm=17)
        _add_para(
            doc,
            f"共计 {len(hist_partners)} 位股东已退出。",
            size=9,
            italic=False,
            source_tag="qcc-history.get_historical_shareholders",
        )

    # §4.4 实控人
    if should_render_section("4.4", snap):
        ac_data = snap.raw.get("mcp__qcc-company__get_actual_controller") or {}
        ac_list = _safe_list(ac_data, "实际控制人信息")
        bo_data = snap.raw.get("mcp__qcc-company__get_beneficial_owners") or {}
        bo_list = _safe_list(bo_data, "受益人信息") or _safe_list(bo_data, "受益所有人信息")

        if ac_list or bo_list:
            _add_heading(doc, "4.4 实际控制人认定", level=2)
            if ac_list:
                rows = [
                    [
                        a.get("实际控制人名称") or a.get("姓名") or "",
                        a.get("直接持股比例") or "",
                        a.get("总持股比例") or "",
                        a.get("表决权比例") or "",
                    ]
                    for a in ac_list
                ]
                _add_table(doc, ["姓名", "直接持股", "总持股", "表决权"], rows)
            if bo_list:
                _add_para(doc, "受益所有人：", bold=True)
                rows = [
                    [
                        o.get("姓名") or o.get("受益人名称") or "",
                        o.get("受益股份比例") or o.get("持股比例") or "",
                        o.get("受益类型") or "",
                    ]
                    for o in bo_list
                ]
                _add_table(doc, ["姓名", "受益比例", "类型"], rows)
            _add_para(
                doc,
                "",
                source_tag="qcc-company.get_actual_controller + get_beneficial_owners",
            )


def render_s5_operations(doc: Document, snap: CompanySnapshot) -> None:
    br_data = snap.raw.get("mcp__qcc-company__get_branches") or {}
    inv_data = snap.raw.get("mcp__qcc-company__get_external_investments") or {}
    lic_data = snap.raw.get("mcp__qcc-operation__get_administrative_license") or {}

    has_content = (
        _safe_list(br_data, "分支机构信息")
        or _safe_list(inv_data, "对外投资信息")
        or _safe_list(lic_data, "行政许可信息")
        or _safe_list(lic_data, "行政许可")
    )
    if not has_content and not any(
        should_render_section(s, snap) for s in ("5.4", "5.5", "5.6")
    ):
        return

    _add_heading(doc, "五、经营轨迹", level=1)

    br_list = _safe_list(br_data, "分支机构信息")
    if br_list:
        _add_heading(doc, "5.2 分支机构", level=2)
        rows = [
            [
                b.get("分支机构名称") or b.get("企业名称") or "",
                b.get("负责人") or "",
                b.get("成立日期") or "",
                b.get("登记状态") or b.get("状态") or "",
            ]
            for b in br_list
        ]
        _add_table(doc, ["分支名称", "负责人", "成立日期", "状态"], rows)
        _add_para(doc, "", source_tag="qcc-company.get_branches")

    inv_list = _safe_list(inv_data, "对外投资信息")
    if inv_list:
        _add_heading(doc, "5.2 对外投资（子/参股）", level=2)
        rows = [
            [
                i.get("被投资企业名称") or i.get("企业名称") or "",
                i.get("持股比例") or "",
                i.get("成立日期") or "",
                i.get("登记状态") or i.get("状态") or "",
            ]
            for i in inv_list
        ]
        _add_table(doc, ["被投资企业", "持股比例", "成立日期", "状态"], rows)
        _add_para(doc, "", source_tag="qcc-company.get_external_investments")

    lic_list = _safe_list(lic_data, "行政许可信息") or _safe_list(lic_data, "行政许可")
    if lic_list:
        _add_heading(doc, "5.2 行政许可", level=2)
        rows = [
            [
                l.get("许可文件名称") or l.get("许可内容") or "",
                l.get("许可决定日期") or l.get("起始日期") or "",
                l.get("有效期至") or l.get("终止日期") or "",
                l.get("许可机关") or l.get("发证机关") or "",
            ]
            for l in lic_list
        ]
        _add_table(doc, ["许可名称", "起始", "终止", "发证机关"], rows)
        _add_para(doc, "", source_tag="qcc-operation.get_administrative_license")

    # §5.4 董监高
    if should_render_section("5.4", snap):
        kp_data = snap.raw.get("mcp__qcc-company__get_key_personnel") or {}
        employees = _safe_list(kp_data, "主要人员信息")
        hist_exec_data = snap.raw.get("mcp__qcc-history__get_historical_executives") or {}
        hist_execs = _safe_list(hist_exec_data, "历史主要人员信息")
        if employees or hist_execs:
            _add_heading(doc, "5.4 董监高与治理", level=2)
            if employees:
                _add_para(doc, "现任董监高：", bold=True)
                rows = [[e.get("姓名") or "", e.get("职务") or "", e.get("持股比例") or ""] for e in employees]
                _add_table(doc, ["姓名", "职务", "持股比例"], rows)
                _add_para(doc, "", source_tag="qcc-company.get_key_personnel")
            if hist_execs:
                _add_para(doc, "历任董监高时间线：", bold=True)
                png = build_timeline.timeline_historical_executives(hist_execs)
                _embed_timeline_png(doc, png, width_cm=16)
                _add_para(doc, "", size=9, source_tag="qcc-history.get_historical_executives")

    # §5.5 关联方
    if should_render_section("5.5", snap):
        render_section_55_associates(doc, snap)

    # §5.6 风险
    if should_render_section("5.6", snap):
        _add_heading(doc, "5.6 处罚与担保记录", level=2)
        _render_risk_tables(doc, snap)


def render_section_55_associates(doc: Document, snap: CompanySnapshot) -> None:
    top_n = TOP_N_BY_TIER.get(snap.tier, 5)

    person_data_map: dict[str, dict] = {}
    for tool_name in [
        "mcp__qcc-executive__get_personnel_positions",
        "mcp__qcc-executive__get_personnel_controlled_companies",
        "mcp__qcc-executive__get_personnel_investments",
    ]:
        bucket = snap.raw.get(tool_name)
        if not isinstance(bucket, list):
            continue
        for entry in bucket:
            person = entry.get("person") or ""
            result = entry.get("result") or {}
            if not person or _is_empty_response(result):
                continue
            person_data_map.setdefault(person, {})[tool_name] = result

    if not person_data_map:
        return

    _add_heading(doc, "5.5 关联方网络（公开可查）", level=2)
    _add_para(
        doc,
        f"按档位 {snap.tier}，间接投资展示 Top {top_n}；直接投资与兼职任职全量呈现。",
        italic=False,
        size=9,
    )

    for person, tools in person_data_map.items():
        _add_heading(doc, f"关键人：{person}", level=3)

        pos_data = tools.get("mcp__qcc-executive__get_personnel_positions") or {}
        positions = _safe_list(pos_data, "董监高-在外任职信息")
        if positions:
            _add_para(doc, "在外兼职：", bold=True)
            rows = [
                [p.get("企业名称") or "", p.get("职位") or "", p.get("状态") or "",
                 p.get("地区") or "", p.get("行业") or ""]
                for p in positions
            ]
            _add_table(doc, ["企业", "职位", "状态", "地区", "行业"], rows)
            _add_para(doc, "", source_tag="qcc-executive.get_personnel_positions")

        ctrl_data = tools.get("mcp__qcc-executive__get_personnel_controlled_companies") or {}
        controlled = _safe_list(ctrl_data, "董监高-控制企业信息")
        if controlled:
            _add_para(doc, f"直接控制企业（共 {len(controlled)} 家）：", bold=True)
            rows = [[c.get("企业名称") or "", c.get("投资比例") or "",
                     c.get("状态") or "", c.get("所属行业") or ""] for c in controlled]
            _add_table(doc, ["企业", "投资比例", "状态", "行业"], rows)
            _add_para(doc, "", source_tag="qcc-executive.get_personnel_controlled_companies")

        inv_data = tools.get("mcp__qcc-executive__get_personnel_investments") or {}
        direct_inv = _safe_list(inv_data, "直接对外投资")
        indirect_inv = _safe_list(inv_data, "间接对外投资")

        if direct_inv:
            _add_para(doc, f"直接投资（共 {len(direct_inv)} 家，全量展示）：", bold=True)
            rows = [[d.get("企业名称") or "", d.get("持股比例") or "",
                     d.get("状态") or "", d.get("地区") or ""] for d in direct_inv]
            _add_table(doc, ["企业", "持股", "状态", "地区"], rows)
            _add_para(doc, "", source_tag="qcc-executive.get_personnel_investments.直接对外投资")

        if indirect_inv:
            truncated = indirect_inv[:top_n]
            suffix = f"，展示 Top {len(truncated)}" if len(indirect_inv) > top_n else ""
            _add_para(doc, f"间接投资（共 {len(indirect_inv)} 家{suffix}）：", bold=True)
            rows = [[i.get("被间接投资企业名称") or "", i.get("间接持股比例") or "",
                     i.get("状态") or "", i.get("地区") or ""] for i in truncated]
            _add_table(doc, ["被投资企业", "间接持股", "状态", "地区"], rows)
            if len(indirect_inv) > top_n:
                _add_para(
                    doc,
                    f"（共 {len(indirect_inv)} 家间接投资，仅展示前 {top_n} 名）",
                    italic=False, size=8,
                )
            _add_para(doc, "", source_tag="qcc-executive.get_personnel_investments.间接对外投资")


def _render_risk_tables(doc: Document, snap: CompanySnapshot) -> None:
    risk_map = [
        ("mcp__qcc-risk__get_administrative_penalty", "行政处罚", "行政处罚信息",
         ["处罚文号", "违法行为类型", "处罚机关", "处罚决定日期"]),
        ("mcp__qcc-risk__get_guarantee_info", "对外担保", "对外担保信息",
         ["债权人", "债务人", "担保金额", "担保期限"]),
        ("mcp__qcc-risk__get_equity_pledge_info", "股权质押", "股权质押信息",
         ["出质人", "质权人", "出质股权数额", "登记日期"]),
        ("mcp__qcc-risk__get_equity_freeze", "股权冻结", "股权冻结信息",
         ["被执行人", "冻结金额", "执行法院", "冻结起始日期"]),
        ("mcp__qcc-risk__get_dishonest_info", "失信被执行人", "失信被执行人信息",
         ["案号", "执行法院", "执行内容", "立案日期"]),
    ]
    for tool, title, list_key, fields in risk_map:
        data = snap.raw.get(tool)
        if _is_empty_response(data):
            continue
        items = _safe_list(data, list_key) or _safe_list(data, "信息")
        if not items:
            continue
        _add_para(doc, title, bold=True)
        rows = [[item.get(f) or "" for f in fields] for item in items]
        _add_table(doc, fields, rows)
        _add_para(doc, "", source_tag=tool.replace("mcp__", ""))


def render_s6_address(doc: Document, snap: CompanySnapshot) -> None:
    hist_reg = snap.raw.get("mcp__qcc-history__get_historical_registration") or {}
    addr_list = _safe_list(hist_reg, "历史地址列表")
    current_addr = snap.profile.get("注册地址") or snap.profile.get("地址")

    if not (addr_list or current_addr):
        return

    _add_heading(doc, "六、地址变迁", level=1)

    if current_addr:
        _add_para(doc, f"当前注册地址：{current_addr}", source_tag="qcc-company.get_company_profile")

    if addr_list:
        png = build_timeline.timeline_address(addr_list)
        _embed_timeline_png(doc, png, width_cm=17)
        _add_para(doc, "", size=9, italic=False, source_tag="qcc-history.get_historical_registration")


# ---------------------------------------------------------------------------
# §7 总结 · Decision Pack 风格卡片
# ---------------------------------------------------------------------------


def render_s7_summary(doc: Document, snap: CompanySnapshot) -> None:
    _add_heading(doc, "七、总结", level=1)
    rating = _compute_overall_rating(snap)

    # Decision Pack 摘要卡：3 行 × 2 列
    t = doc.add_table(rows=4, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        ("整体评价", rating["label"]),
        ("股权清洁度", rating["equity"]),
        ("经营稳定度", rating["operation"]),
        ("成长轨迹", rating["growth"]),
    ]
    for i, (k, v) in enumerate(rows):
        c0 = t.rows[i].cells[0]
        c0.text = k
        c1 = t.rows[i].cells[1]
        c1.text = v
        for cell in (c0, c1):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Microsoft YaHei"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(10)
                    if cell is c0:
                        run.bold = True

    _add_para(doc, "")
    _add_para(doc, rating["narrative"], size=10)

    if should_render_section("7.fin", snap):
        fi = snap.raw.get("mcp__qcc-company__get_financial_data") or {}
        if not _is_empty_response(fi):
            _add_heading(doc, "7.X 财务摘要（近 3 年对比）", level=2)
            _render_financial_multi_year(doc, fi)


def _compute_overall_rating(snap: CompanySnapshot) -> dict[str, str]:
    raw = snap.raw

    freeze = raw.get("mcp__qcc-risk__get_equity_freeze") or {}
    pledge = raw.get("mcp__qcc-risk__get_equity_pledge_info") or {}
    has_freeze = not _is_empty_response(freeze)
    has_pledge = not _is_empty_response(pledge)
    equity = "存在冻结/质押，需核查" if (has_freeze or has_pledge) else "未发现公开冻结/质押记录"

    biz_exc = raw.get("mcp__qcc-risk__get_business_exception") or {}
    tax_abn = raw.get("mcp__qcc-risk__get_tax_abnormal") or {}
    penalty = raw.get("mcp__qcc-risk__get_administrative_penalty") or {}
    has_neg = any(not _is_empty_response(x) for x in (biz_exc, tax_abn, penalty))
    operation = "存在经营/税务/处罚异常" if has_neg else "未发现公开经营异常"

    changes = raw.get("mcp__qcc-company__get_change_records") or {}
    change_count = len(_safe_list(changes, "变更记录信息"))
    if change_count > 30:
        growth = f"变更频繁（{change_count} 次），建议审视变动动机"
    elif change_count > 10:
        growth = f"变更较多（{change_count} 次），属成熟期常态"
    else:
        growth = f"变更较少（{change_count} 次），属稳定经营"

    neg_flags = sum([has_freeze, not _is_empty_response(biz_exc), not _is_empty_response(penalty)])
    if neg_flags >= 2:
        label = "evasive_high_risk 高风险 / 待核查"
        narr = "公开记录中存在多项异常信号，建议线下尽调补强。"
    elif neg_flags == 1:
        label = "recovered 历史有瑕疵 / 需核查"
        narr = "公开记录中存在个别异常信号，不排除已处置。"
    else:
        label = "clean_baseline 清洁基线"
        narr = "公开记录未发现显著风险事件。"

    return {"label": label, "narrative": narr, "equity": equity, "operation": operation, "growth": growth}


def _render_financial_multi_year(doc: Document, fi: dict) -> None:
    fin_data = _safe_list(fi, "财务数据信息")
    if not fin_data:
        return

    fin_data = sorted(fin_data, key=lambda x: x.get("报告期") or "", reverse=True)[:3]
    if not fin_data:
        return

    periods = [f.get("报告期") or f"Y{i}" for i, f in enumerate(fin_data)]

    def collect(path: list[str]) -> dict[str, list[str]]:
        out: dict[str, list[str]] = {}
        all_keys: list[str] = []
        for f in fin_data:
            d: Any = f.get("指标详情") or {}
            for p in path:
                d = d.get(p) if isinstance(d, dict) else None
                if d is None:
                    break
            if isinstance(d, dict):
                for k in d:
                    if k not in all_keys:
                        all_keys.append(k)
        for f in fin_data:
            d = f.get("指标详情") or {}
            for p in path:
                d = d.get(p) if isinstance(d, dict) else {}
            if not isinstance(d, dict):
                d = {}
            for k in all_keys:
                out.setdefault(k, []).append(str(d.get(k, "")))
        return out

    sections = [
        ("主要财务指标", ["主要财务指标"]),
        ("利润表", ["财务报表", "利润表"]),
        ("资产负债表", ["财务报表", "资产负债表"]),
        ("现金流量表", ["财务报表", "现金流量表"]),
        ("盈利能力", ["分析数据", "盈利能力"]),
        ("偿还能力", ["分析数据", "偿还能力"]),
        ("成长能力", ["分析数据", "成长能力"]),
    ]
    for title, path in sections:
        data = collect(path)
        if not data:
            continue
        _add_para(doc, title, bold=True, size=11)
        rows = [[k] + vals for k, vals in data.items()]
        _add_table(doc, ["指标"] + periods, rows)

    _add_para(doc, "", source_tag="qcc-company.get_financial_data")


def render_s0_milestones(doc: Document, snap: CompanySnapshot) -> None:
    """§0 里程碑总览（V1.1 · PPT P5 风格上升曲线图）
    聚合：成立 + 改制 + 融资节点 + 重大荣誉 + 上榜 + 大额认证
    """
    events: list[dict] = []

    # 1. 成立
    reg = snap.raw.get("mcp__qcc-company__get_company_registration_info") or {}
    start_date = reg.get("成立日期") or ""
    if start_date:
        events.append({"year": start_date[:4], "title": "企业成立", "sub": ""})

    # 2. 历史名称变更（关键节点，如改制）
    hist_reg = snap.raw.get("mcp__qcc-history__get_historical_registration") or {}
    for nm in _safe_list(hist_reg, "历史名称列表"):
        end = nm.get("终止日期") or ""
        if end:
            events.append({
                "year": end[:4],
                "title": "改制/更名",
                "sub": f"更名为{nm.get('历史名称', '')[:8]}",
            })

    # 3. 融资节点（每年取 1 条最重要）
    financing = snap.financing or {}
    equity_fin = financing.get("股权融资") or {}
    vc_rounds = equity_fin.get("创投融资") or []
    fin_by_year: dict[str, dict] = {}
    for r in vc_rounds:
        y = (r.get("融资日期") or "")[:4]
        if not y:
            continue
        # 每年只取金额最大或最新一轮
        if y not in fin_by_year:
            fin_by_year[y] = r
    for y, r in fin_by_year.items():
        rnd = r.get("融资轮次") or ""
        amt = r.get("融资金额") or ""
        events.append({
            "year": y,
            "title": f"获{rnd}融资",
            "sub": amt if amt and amt != "未披露" else "",
        })

    # 4. 重大荣誉/认证（国家级 + 关键认证）
    honor_data = snap.raw.get("mcp__qcc-operation__get_honor_info") or {}
    important_honors = []
    for h in _safe_list(honor_data, "荣誉信息"):
        level = h.get("级别") or ""
        name = h.get("名称") or ""
        year = h.get("认证年份") or ""
        # 筛选国家级/关键认证
        if ("国家级" in level) or any(kw in name for kw in ["高新技术", "独角兽", "DCMM", "央行", "示范"]):
            important_honors.append({"year": year, "name": name})
    # 按年聚合，每年一条
    honor_by_year: dict[str, str] = {}
    for h in important_honors:
        y, n = h["year"], h["name"]
        if y and y not in honor_by_year:
            honor_by_year[y] = n
    for y, n in honor_by_year.items():
        events.append({"year": y, "title": "重要认证/荣誉", "sub": n[:12]})

    # 5. 上市/挂牌（如有）
    listing = snap.listing or {}
    if not _is_empty_response(listing):
        status = listing.get("上市状态") or ""
        date = listing.get("上市日期") or ""
        if date:
            events.append({"year": date[:4], "title": "上市/挂牌",
                          "sub": listing.get("股票简称", "")[:8]})

    # 按年份去重（每年 1-2 条），排序
    year_group: dict[str, list] = {}
    for e in events:
        y = e["year"]
        if y:
            year_group.setdefault(y, []).append(e)

    # 每年合并同年事件
    merged: list[dict] = []
    for y in sorted(year_group.keys()):
        items = year_group[y]
        if len(items) == 1:
            merged.append(items[0])
        else:
            # 合并多个事件到一个节点
            titles = "、".join(sorted({i["title"] for i in items}))
            subs = "、".join(filter(None, sorted({i["sub"] for i in items if i.get("sub")})))
            merged.append({"year": y, "title": titles, "sub": subs[:20]})

    if len(merged) < 2:
        return  # 节点太少不画

    _add_heading(doc, "发展里程碑", level=1)
    _add_para(
        doc,
        f"{snap.profile.get('企业名称') or snap.keyword} 自 {merged[0]['year']} 年"
        f"至今的重要发展里程碑，共 {len(merged)} 个关键节点。",
        size=10, italic=False,
    )
    # V1.1.1：回退为圆点+竖线时间轴（与其他章节风格一致）
    timeline_items = [
        {
            "date": m["year"],
            "content": m["title"],
            "sub": m.get("sub", ""),
            "known": True,
        }
        for m in merged
    ]
    png = build_timeline.render_single_column_timeline(
        timeline_items, width=14.0, row_h=0.7
    )
    _embed_timeline_png(doc, png, width_cm=16)
    _add_para(
        doc, "",
        size=9, italic=False,
        source_tag="多源聚合：registration / financing / honor / listing",
    )


def render_s8_development_overview(doc: Document, snap: CompanySnapshot) -> None:
    """§8 发展综览（V1.1 · 立体维度聚合）"""
    # 检查是否有足够数据启用本章
    honor = snap.raw.get("mcp__qcc-operation__get_honor_info") or {}
    ranking = snap.raw.get("mcp__qcc-operation__get_ranking_list_info") or {}
    quals = snap.raw.get("mcp__qcc-operation__get_qualifications") or {}
    patent = snap.raw.get("mcp__qcc-ipr__get_patent_info") or {}
    trademark = snap.raw.get("mcp__qcc-ipr__get_trademark_info") or {}
    sw_copyright = snap.raw.get("mcp__qcc-ipr__get_software_copyright_info") or {}
    work_copyright = snap.raw.get("mcp__qcc-ipr__get_copyright_work_info") or {}
    internet = snap.raw.get("mcp__qcc-ipr__get_internet_service_info") or {}
    annual = snap.raw.get("mcp__qcc-company__get_annual_reports") or {}
    recruit = snap.raw.get("mcp__qcc-operation__get_recruitment_info") or {}
    credit = snap.raw.get("mcp__qcc-operation__get_credit_evaluation") or {}

    has_any = any(
        not _is_empty_response(x)
        for x in (honor, ranking, quals, patent, trademark, sw_copyright,
                  work_copyright, internet, annual, recruit, credit)
    )
    if not has_any:
        return

    _add_heading(doc, "八、发展综览", level=1)
    _add_para(
        doc,
        "从 资质许可 / 知识产权 / 业务扩张 / 行业地位 四个维度立体呈现企业发展。",
        size=10, italic=False,
    )

    # §8.1 许可资质矩阵
    quals_list = _safe_list(quals, "资质证书信息")
    if quals_list:
        _add_heading(doc, "8.1 许可与资质", level=2)
        # 按资质名称分类计数
        cat_count: dict[str, int] = {}
        active_count = 0
        for q in quals_list:
            cat = q.get("资质名称") or "其他"
            cat_count[cat] = cat_count.get(cat, 0) + 1
            if q.get("证书状态") == "有效":
                active_count += 1
        _add_para(
            doc,
            f"共 {len(quals_list)} 项资质证书，其中有效 {active_count} 项。",
            size=10,
        )
        # 列出前 8 条有效资质
        active_quals = [q for q in quals_list if q.get("证书状态") == "有效"][:8]
        if active_quals:
            rows = [
                [
                    q.get("资质名称") or "",
                    q.get("类别与等级") or "—",
                    q.get("发证日期") or "—",
                    q.get("有效期至") or "长期有效",
                ]
                for q in active_quals
            ]
            _add_table(doc, ["资质名称", "类别/等级", "发证日期", "有效期至"], rows)
        _add_para(doc, "", size=9, source_tag="qcc-operation.get_qualifications")

    # §8.2 知识产权布局
    ipr_has_data = any(
        not _is_empty_response(x)
        for x in (patent, trademark, sw_copyright, work_copyright, internet)
    )
    if ipr_has_data:
        _add_heading(doc, "8.2 知识产权布局", level=2)

        # V1.1.1：统一用计数表（与 kyb-opus 风格一致）
        patent_list = _safe_list(patent, "专利信息")
        tm_list = _safe_list(trademark, "商标信息")
        sw_list = _safe_list(sw_copyright, "软件著作权信息")
        wk_list = _safe_list(work_copyright, "作品著作权信息")
        isp_list = _safe_list(internet, "ICP备案信息") or _safe_list(internet, "互联网备案信息")

        count_rows = []
        if patent_list:
            invention_cnt = sum(1 for p in patent_list if "发明" in (p.get("专利类型") or ""))
            design_cnt = sum(1 for p in patent_list if "外观" in (p.get("专利类型") or ""))
            count_rows.append(["专利", str(len(patent_list)),
                               f"发明 {invention_cnt} / 外观 {design_cnt}"])
        if tm_list:
            count_rows.append(["商标", str(len(tm_list)), "—"])
        if sw_list:
            count_rows.append(["软件著作权", str(len(sw_list)), "—"])
        if wk_list:
            count_rows.append(["作品著作权", str(len(wk_list)), "—"])
        if isp_list:
            count_rows.append(["互联网备案", str(len(isp_list)), "ICP/APP/小程序/算法"])
        if count_rows:
            _add_para(doc, "知识产权数量概览：", bold=True)
            _add_table(doc, ["类型", "数量", "细分说明"], count_rows)

        # V1.1.1：专利年度分布改为圆点+竖线时间轴
        if patent_list and len(patent_list) >= 3:
            patent_year: dict[str, int] = {}
            for p in patent_list:
                y = (p.get("申请日期") or "")[:4]
                if y:
                    patent_year[y] = patent_year.get(y, 0) + 1
            if len(patent_year) >= 2:
                _add_para(doc, "专利申请年度分布：", bold=True)
                yearly_items = [
                    {"date": y, "content": f"{patent_year[y]} 项", "sub": "", "known": True}
                    for y in sorted(patent_year.keys(), reverse=True)
                ]
                png = build_timeline.render_single_column_timeline(
                    yearly_items, width=11.0, row_h=0.55
                )
                _embed_timeline_png(doc, png, width_cm=13)

        _add_para(doc, "", size=9, source_tag="qcc-ipr.patent/trademark/copyright")

    # §8.3 业务扩张（分支/对外投资/招聘/员工规模）
    branches = snap.raw.get("mcp__qcc-company__get_branches") or {}
    br_list = _safe_list(branches, "分支机构信息")
    inv = snap.raw.get("mcp__qcc-company__get_external_investments") or {}
    inv_list = _safe_list(inv, "对外投资信息")

    # 员工规模（从年报社保人数抽）
    annual_list = _safe_list(annual, "企业年报信息")
    staff_by_year: dict[str, int] = {}
    for ar in annual_list:
        y = (ar.get("年报年度") or "")[:4]
        social = ar.get("社保信息") or {}
        staff_str = social.get("城镇职工基本养老保险") or ""
        if "人" in staff_str:
            try:
                n = int(staff_str.replace("人", "").strip())
                if y:
                    staff_by_year[y] = n
            except Exception:
                pass

    if br_list or inv_list or staff_by_year:
        _add_heading(doc, "8.3 业务扩张与团队规模", level=2)

        # V1.1.1：员工规模改为圆点+竖线时间轴
        if staff_by_year:
            _add_para(doc, "员工规模年度趋势（社保参保人数）：", bold=True)
            staff_items = [
                {"date": y, "content": f"{staff_by_year[y]} 人",
                 "sub": "", "known": True}
                for y in sorted(staff_by_year.keys(), reverse=True)
            ]
            png = build_timeline.render_single_column_timeline(
                staff_items, width=10.0, row_h=0.55
            )
            _embed_timeline_png(doc, png, width_cm=12)

        if br_list:
            _add_para(doc, f"分支机构（{len(br_list)} 家）：", bold=True)
            rows = [
                [b.get("分支机构名称") or b.get("企业名称") or "",
                 b.get("地区") or "—",
                 b.get("成立日期") or "—",
                 b.get("登记状态") or b.get("状态") or ""]
                for b in br_list[:10]
            ]
            _add_table(doc, ["分支名称", "地区", "成立日期", "状态"], rows)

        if inv_list:
            _add_para(doc, f"对外投资（{len(inv_list)} 家）：", bold=True)
            rows = [
                [i.get("被投资企业名称") or i.get("企业名称") or "",
                 i.get("持股比例") or "—",
                 i.get("成立日期") or "—",
                 i.get("登记状态") or i.get("状态") or ""]
                for i in inv_list[:10]
            ]
            _add_table(doc, ["被投资企业", "持股比例", "成立日期", "状态"], rows)

        _add_para(doc, "", size=9, source_tag="qcc-company.get_branches + get_external_investments + get_annual_reports")

    # §8.4 行业地位与荣誉
    honor_list = _safe_list(honor, "荣誉信息")
    ranking_list = _safe_list(ranking, "上榜榜单信息")
    has_honor = bool(honor_list or ranking_list)

    if has_honor:
        _add_heading(doc, "8.4 行业地位与荣誉", level=2)

        # V1.1.1：统一用计数表（与 kyb-opus 风格一致）
        count_rows = []
        if honor_list:
            national = sum(1 for h in honor_list if "国家级" in (h.get("级别") or ""))
            prov = sum(1 for h in honor_list if "省级" in (h.get("级别") or ""))
            count_rows.append([
                "企业荣誉", str(len(honor_list)),
                f"其中国家级 {national} 项 / 省级 {prov} 项",
            ])
        if ranking_list:
            count_rows.append(["上榜榜单", str(len(ranking_list)), "含互联网百强 / 独角兽 / AI 榜单等"])
        patent_cnt = len(_safe_list(patent, "专利信息"))
        if patent_cnt:
            count_rows.append(["专利信息", str(patent_cnt), "—"])

        if count_rows:
            _add_para(doc, "行业地位与荣誉概览：", bold=True)
            _add_table(doc, ["维度", "数量", "说明"], count_rows)

        # 重点荣誉列表（国家级 + 省级前 8 条）
        if honor_list:
            key_honors = [
                h for h in honor_list
                if any(kw in (h.get("级别") or "") for kw in ["国家级", "省级"])
            ][:10]
            if key_honors:
                _add_para(doc, "重点荣誉（按授予年份）：", bold=True)
                rows = [
                    [h.get("认证年份") or "",
                     h.get("名称") or "",
                     h.get("级别") or "",
                     h.get("发布单位") or ""]
                    for h in key_honors
                ]
                _add_table(doc, ["年份", "荣誉名称", "级别", "颁发单位"], rows)

        # 上榜榜单 Top 8
        if ranking_list:
            top_rank = ranking_list[:8]
            _add_para(doc, "近期上榜榜单（Top 8）：", bold=True)
            rows = [
                [r.get("发布日期") or "",
                 r.get("榜单名称") or "",
                 r.get("榜内排名") or "—",
                 r.get("来源") or ""]
                for r in top_rank
            ]
            _add_table(doc, ["发布日期", "榜单名称", "排名", "来源"], rows)

        _add_para(doc, "", size=9, source_tag="qcc-operation.get_honor_info + get_ranking_list_info")

    # §8.5 信用评级（如有）
    if not _is_empty_response(credit):
        credit_list = _safe_list(credit, "信用评价信息") or _safe_list(credit, "信用评级信息")
        if credit_list:
            _add_heading(doc, "8.5 官方信用评级", level=2)
            rows = [
                [c.get("评价类型") or c.get("评级类型") or "",
                 c.get("等级") or c.get("信用等级") or "",
                 c.get("评价单位") or c.get("发证机关") or "",
                 c.get("评价年度") or c.get("评定年度") or ""]
                for c in credit_list
            ]
            _add_table(doc, ["评级类型", "等级", "评定单位", "评价年度"], rows)
            _add_para(doc, "", size=9, source_tag="qcc-operation.get_credit_evaluation")


def render_appendix_timeline(doc: Document, snap: CompanySnapshot) -> None:
    cr_data = snap.raw.get("mcp__qcc-company__get_change_records") or {}
    changes = _safe_list(cr_data, "变更记录信息")
    if not changes:
        return

    doc.add_page_break()
    _add_heading(doc, "附录 · 工商变更大事年表", level=1)
    _add_para(
        doc,
        f"共 {len(changes)} 条变更记录（按时间倒序，最多展示 40 条）。",
        italic=False, size=9,
    )
    png = build_timeline.timeline_change_records(changes, max_items=40)
    _embed_timeline_png(doc, png, width_cm=17)
    _add_para(doc, "", size=9, source_tag="qcc-company.get_change_records")


def render_footer(doc: Document, snap: CompanySnapshot) -> None:
    if snap.uncollected:
        _add_heading(doc, "附：未采集字段", level=2)
        for t in snap.uncollected:
            _add_para(doc, f"· {t}", size=9)

    _add_heading(doc, "本报告数据源", level=2)
    _add_para(
        doc,
        f"数据源：企查查 MCP 146 工具（V1.1.3，纯 MCP 驱动 + 时间轴 PNG 嵌入）。总成本："
        f"{snap.cost_summary.get('total_cent', 0)}c "
        f"(¥{snap.cost_summary.get('total_rmb', 0)})。",
        size=9, italic=False,
    )
    _add_para(
        doc,
        "本报告仅依赖企查查 MCP 可公开查询字段；非公示数据（代持定性、一致行动、独立性评估、"
        "关联交易明细等）需客户 V2.0 自扩展。",
        size=9, italic=False,
    )


def build(snapshot: CompanySnapshot, output_path: str) -> str:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    _set_default_font(doc)

    render_title(doc, snapshot)
    render_s1_overview(doc, snapshot)
    render_s0_milestones(doc, snapshot)  # V1.1：里程碑总览（放 §1 之后）
    render_s2_capital(doc, snapshot)
    render_s3_name(doc, snapshot)
    render_s4_equity(doc, snapshot)
    render_s5_operations(doc, snapshot)
    render_s6_address(doc, snapshot)
    render_s7_summary(doc, snapshot)
    render_s8_development_overview(doc, snapshot)  # V1.1：发展综览
    render_appendix_timeline(doc, snapshot)
    render_footer(doc, snapshot)

    doc.save(output_path)
    return output_path


def build_from_json(json_path: str, output_path: str) -> str:
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    snap = CompanySnapshot(
        keyword=data.get("keyword", ""),
        tier=data.get("tier", TIER_MICRO),
        tier_reason=data.get("tier_reason", ""),
        profile=data.get("profile") or {},
        listing=data.get("listing") or {},
        financing=data.get("financing") or {},
        raw=data.get("raw") or {},
        cost_summary=data.get("cost_summary") or {},
        uncollected=data.get("uncollected") or [],
    )
    return build(snap, output_path)


def convert_to_pdf(docx_path: str) -> str:
    import subprocess
    out_dir = os.path.dirname(os.path.abspath(docx_path)) or "."
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        check=True,
    )
    return docx_path.replace(".docx", ".pdf")


def main() -> None:
    ap = argparse.ArgumentParser(description="企业历史沿革报告生成器（V1.1.3）")
    ap.add_argument("--input", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--pdf", action="store_true")
    args = ap.parse_args()

    docx_path = build_from_json(args.input, args.output)
    print(f"DOCX 已生成：{docx_path}")
    if args.pdf:
        pdf_path = convert_to_pdf(docx_path)
        print(f"PDF 已生成：{pdf_path}")


if __name__ == "__main__":
    main()
